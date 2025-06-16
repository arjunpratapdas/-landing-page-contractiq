import os
import logging
from pathlib import Path
from typing import Optional

# Document processing libraries
try:
    import fitz  # PyMuPDF for PDF processing
except ImportError:
    fitz = None

try:
    from docx import Document  # python-docx for DOCX files
except ImportError:
    Document = None

try:
    import pdfplumber  # Alternative PDF processor
except ImportError:
    pdfplumber = None

logger = logging.getLogger(__name__)

class DocumentParser:
    """
    Document parser for extracting text from various file formats
    """
    
    def __init__(self):
        self.supported_formats = {'.pdf', '.doc', '.docx'}
        logger.info("DocumentParser initialized")
    
    def extract_text(self, file_path: str, file_extension: str) -> str:
        """
        Extract text from document based on file extension
        
        Args:
            file_path: Path to the document file
            file_extension: File extension (.pdf, .doc, .docx)
            
        Returns:
            Extracted text content
        """
        try:
            if file_extension.lower() == '.pdf':
                return self._extract_pdf_text(file_path)
            elif file_extension.lower() in ['.doc', '.docx']:
                return self._extract_docx_text(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
        
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """
        Extract text from PDF using PyMuPDF (primary) or pdfplumber (fallback)
        """
        text = ""
        
        # Try PyMuPDF first
        if fitz:
            try:
                doc = fitz.open(file_path)
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    text += page.get_text()
                doc.close()
                logger.info(f"Successfully extracted text from PDF using PyMuPDF: {len(text)} characters")
                return text
            except Exception as e:
                logger.warning(f"PyMuPDF failed: {str(e)}, trying pdfplumber")
        
        # Fallback to pdfplumber
        if pdfplumber:
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                logger.info(f"Successfully extracted text from PDF using pdfplumber: {len(text)} characters")
                return text
            except Exception as e:
                logger.error(f"pdfplumber also failed: {str(e)}")
        
        # If both methods fail
        if not text:
            raise Exception("Could not extract text from PDF. Please ensure the PDF contains readable text.")
        
        return text
    
    def _extract_docx_text(self, file_path: str) -> str:
        """
        Extract text from DOC/DOCX files using python-docx
        """
        if not Document:
            raise Exception("python-docx not available. Cannot process DOC/DOCX files.")
        
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            logger.info(f"Successfully extracted text from DOCX: {len(text)} characters")
            return text
        
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise Exception(f"Could not extract text from document: {str(e)}")
    
    def validate_file(self, file_path: str, max_size_mb: int = 10) -> bool:
        """
        Validate file size and format
        
        Args:
            file_path: Path to file
            max_size_mb: Maximum file size in MB
            
        Returns:
            True if valid, False otherwise
        """
        try:
            # Check if file exists
            if not os.path.exists(file_path):
                logger.error(f"File does not exist: {file_path}")
                return False
            
            # Check file size
            file_size = os.path.getsize(file_path)
            max_size_bytes = max_size_mb * 1024 * 1024
            
            if file_size > max_size_bytes:
                logger.error(f"File too large: {file_size} bytes (max: {max_size_bytes})")
                return False
            
            # Check file extension
            file_extension = Path(file_path).suffix.lower()
            if file_extension not in self.supported_formats:
                logger.error(f"Unsupported file format: {file_extension}")
                return False
            
            return True
        
        except Exception as e:
            logger.error(f"Error validating file: {str(e)}")
            return False
    
    def get_document_metadata(self, file_path: str) -> dict:
        """
        Extract metadata from document
        
        Args:
            file_path: Path to document
            
        Returns:
            Dictionary containing metadata
        """
        metadata = {
            "file_size": os.path.getsize(file_path),
            "file_extension": Path(file_path).suffix.lower(),
            "file_name": Path(file_path).name
        }
        
        try:
            if metadata["file_extension"] == ".pdf" and fitz:
                doc = fitz.open(file_path)
                metadata.update({
                    "page_count": len(doc),
                    "title": doc.metadata.get("title", ""),
                    "author": doc.metadata.get("author", ""),
                    "subject": doc.metadata.get("subject", ""),
                    "creator": doc.metadata.get("creator", "")
                })
                doc.close()
            
            elif metadata["file_extension"] in [".doc", ".docx"] and Document:
                doc = Document(file_path)
                core_props = doc.core_properties
                metadata.update({
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "subject": core_props.subject or "",
                    "created": str(core_props.created) if core_props.created else "",
                    "modified": str(core_props.modified) if core_props.modified else ""
                })
        
        except Exception as e:
            logger.warning(f"Could not extract metadata: {str(e)}")
        
        return metadata